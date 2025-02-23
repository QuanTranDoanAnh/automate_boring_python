#! python3
import docx

doc = docx.Document()
doc.add_paragraph('Hello world!')
paraObject1 = doc.add_paragraph('This is a second paragraph.')
paraObject2 = doc.add_paragraph('This is yet another paragraph.')
paraObject1.add_run(' This text is being added to the second paragraph.')
doc.save('./data/multipleParagraphs.docx')